"""Handlers for importing external profile instructions and documents."""

from __future__ import annotations

import asyncio
import io
import json
import logging
import textwrap
from dataclasses import dataclass, field
from pathlib import Path
from typing import TYPE_CHECKING, Any, Iterable, List, Optional, Sequence, Tuple

from telegram import Update
from telegram.ext import ContextTypes

from app.db import db_rw
from app.utils.fs import ensure_user_dirs, user_dir
from app.utils.ollama import generate
from app.utils.time_utils import operator_now

if TYPE_CHECKING:  # pragma: no cover
    from app.runtime.interfaces import UnifiedWellnessBot

SESSION_KEY = "profile_import_session"
MAX_DOCUMENTS = 10
MAX_TOTAL_CHARACTERS = 120_000
MAX_CONTEXT_CHARACTERS = 70_000
SUMMARY_PREVIEW_CHARS = 420
INLINE_DOC_PREFIX = "pasted_text"


@dataclass
class ImportDocument:
    """In-memory representation of a document collected during bulk import."""

    filename: str
    text: str
    stored_path: Optional[str] = None
    warnings: list[str] = field(default_factory=list)
    source_bytes: Optional[int] = None

    @property
    def char_count(self) -> int:
        return len(self.text or "")


logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Public handlers


async def handle_import_profile(
    bot: "UnifiedWellnessBot",
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Handle `/importprofile` command for quick single-payload imports."""
    if not update.effective_user or not update.message:
        return

    user_id = bot.ensure_user(
        update.effective_user.id,
        update.effective_user.username,
        update.effective_user.first_name,
    )
    if user_id is None:
        await update.message.reply_text(
            "Sorry, I couldn't access your profile. Please try again."
        )
        return

    args = list(getattr(context, "args", []) or [])
    payload = " ".join(args).strip()
    if not payload:
        await update.message.reply_text(
            textwrap.dedent(
                """
                **Profile Import Options:**

                1. Paste text directly:
                   /importprofile {"communication_style": "casual", "topics": ["fitness"]}

                2. Upload a document file (.pdf, .docx, .txt, .md, .json):
                   Just send the document as an attachment and I'll process it automatically!

                3. For multiple documents:
                   Use /importprofilebulk to start a session, then upload files one by one.
                """
            ).strip()
        )
        return

    await update.message.reply_text("📥 Processing your profile import...")

    ensure_user_dirs(int(user_id))
    doc = ImportDocument(
        filename="chatgpt_import.txt",
        text=payload,
        stored_path=_store_text_blob(int(user_id), payload, suffix="single_import"),
    )

    result = await _process_import_documents(bot, int(user_id), [doc])

    summary = _format_user_summary(result)
    bot.invalidate_profile_cache(int(user_id))
    await update.message.reply_text(summary)


async def handle_import_profile_bulk(
    bot: "UnifiedWellnessBot",
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Entry point for `/importprofilebulk` command."""
    if not update.effective_user or not update.message:
        return

    user_id = bot.ensure_user(
        update.effective_user.id,
        update.effective_user.username,
        update.effective_user.first_name,
    )
    if user_id is None:
        await update.message.reply_text(
            "Sorry, I couldn't access your profile. Please try again."
        )
        return

    ensure_user_dirs(int(user_id))
    _start_session(context, int(user_id))

    await update.message.reply_text(
        "📥 **Bulk Profile Import Mode**\n\n"
        "Upload up to 10 documents (txt, json, pdf, docx). "
        "Send them one by one or paste text snippets directly.\n\n"
        "When you're finished, type `done` or `cancel` to exit."
    )


async def handle_cancel_profile_bulk(
    bot: "UnifiedWellnessBot",
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Allow user to cancel the current bulk import session."""
    if not update.message or not update.effective_user:
        return

    session = _get_session(context)
    if not session:
        await update.message.reply_text("No active bulk import session to cancel.")
        return

    _clear_session(context)
    await update.message.reply_text(
        "Bulk import cancelled. No documents were processed."
    )


async def handle_document_upload(
    bot: "UnifiedWellnessBot",
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Handle document uploads during bulk import sessions OR as standalone imports."""
    if not update.message or not update.effective_user or not update.message.document:
        return

    session = _get_session(context)

    # If no active session, handle as a standalone document import
    if not session:
        await _handle_standalone_document_upload(bot, update, context)
        return

    user_id = session["user_id"]
    documents = session["documents"]
    if len(documents) >= MAX_DOCUMENTS:
        await update.message.reply_text(
            f"⚠️ You've reached the limit of {MAX_DOCUMENTS} documents for this import session."
        )
        return

    document = update.message.document
    try:
        file_obj = await context.bot.get_file(document.file_id)
        raw_bytes = bytes(await file_obj.download_as_bytearray())
    except Exception as exc:  # noqa: BLE001
        await update.message.reply_text(
            f"Sorry, I couldn't download {document.file_name}: {exc}"
        )
        return

    filename = document.file_name or "document"
    stored_path = _store_binary_document(user_id, filename, raw_bytes)
    extracted_text, warnings = _extract_text_from_bytes(raw_bytes, filename)

    if not extracted_text.strip():
        warnings.append("No text could be extracted from this document.")

    projected_total = session["total_characters"] + len(extracted_text)
    if projected_total > MAX_TOTAL_CHARACTERS:
        await update.message.reply_text(
            f"⚠️ This document would exceed the {MAX_TOTAL_CHARACTERS:,} character limit for a single import. "
            "Try splitting it into smaller pieces."
        )
        return

    doc_entry = {
        "filename": document.file_name or "document",
        "text": extracted_text,
        "stored_path": stored_path,
        "warnings": warnings,
        "source_bytes": len(raw_bytes),
    }
    documents.append(doc_entry)
    session["total_characters"] = projected_total

    _set_session(context, session)

    warning_line = ""
    if warnings:
        warning_line = "\n⚠️ " + "; ".join(warnings[:2])

    await update.message.reply_text(
        f"✅ Imported **{document.file_name}** ({len(extracted_text):,} chars). "
        f"Total documents: {len(documents)}; total characters: {session['total_characters']:,}.{warning_line}"
    )


async def _handle_standalone_document_upload(
    bot: "UnifiedWellnessBot",
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> None:
    """Handle a single document upload outside of a bulk import session."""
    if not update.message or not update.effective_user or not update.message.document:
        return

    user_id = bot.ensure_user(
        update.effective_user.id,
        update.effective_user.username,
        update.effective_user.first_name,
    )
    if user_id is None:
        await update.message.reply_text(
            "Sorry, I couldn't access your profile. Please try again."
        )
        return

    document = update.message.document
    filename = document.file_name or "document"
    await update.message.reply_text(f"📥 Processing your document: {filename}...")

    try:
        file_obj = await context.bot.get_file(document.file_id)
        raw_bytes = bytes(await file_obj.download_as_bytearray())
    except Exception as exc:  # noqa: BLE001
        await update.message.reply_text(
            f"Sorry, I couldn't download {document.file_name}: {exc}"
        )
        return

    ensure_user_dirs(int(user_id))
    stored_path = _store_binary_document(int(user_id), filename, raw_bytes)
    extracted_text, warnings = _extract_text_from_bytes(raw_bytes, filename)

    if not extracted_text.strip():
        warning_msg = "⚠️ No text could be extracted from this document."
        if warnings:
            warning_msg += f" {'; '.join(warnings)}"
        await update.message.reply_text(warning_msg)
        return

    # Create an ImportDocument and process it
    doc = ImportDocument(
        filename=filename,
        text=extracted_text,
        stored_path=stored_path,
        warnings=warnings,
        source_bytes=len(raw_bytes),
    )

    result = await _process_import_documents(bot, int(user_id), [doc])
    summary = _format_user_summary(result)
    bot.invalidate_profile_cache(int(user_id))

    await update.message.reply_text(summary)


async def maybe_handle_bulk_import_text(
    bot: "UnifiedWellnessBot",
    update: Update,
    context: ContextTypes.DEFAULT_TYPE,
) -> bool:
    """Intercept text messages while a bulk import session is active.

    Returns:
        True if the message was handled here and should not be processed by the
        main conversation handler.
    """
    if not update.message or not update.effective_user:
        return False

    session = _get_session(context)
    if not session:
        return False

    text = update.message.text or ""
    lowered = text.strip().lower()

    if not text.strip():
        await update.message.reply_text(
            "Please provide document content or type `done` to finish."
        )
        return True

    if lowered in {"done", "finished", "complete"}:
        await update.message.reply_text(
            f"⏳ Processing {len(session['documents'])} document(s)... this may take a moment."
        )
        documents = _materialize_documents(session["documents"])
        result = await _process_import_documents(bot, session["user_id"], documents)
        _clear_session(context)
        summary = _format_user_summary(result)
        bot.invalidate_profile_cache(session["user_id"])
        await update.message.reply_text(summary)
        return True

    if lowered in {"cancel", "stop"}:
        _clear_session(context)
        await update.message.reply_text(
            "Bulk import cancelled. No documents were processed."
        )
        return True

    # Treat as inline pasted text
    inline_index = session.setdefault("inline_counter", 0) + 1
    session["inline_counter"] = inline_index
    filename = f"{INLINE_DOC_PREFIX}_{inline_index:02d}.txt"
    projected_total = session["total_characters"] + len(text)
    if projected_total > MAX_TOTAL_CHARACTERS:
        await update.message.reply_text(
            f"⚠️ Adding this text would exceed the {MAX_TOTAL_CHARACTERS:,} character limit. "
            "Consider removing some content or splitting into multiple imports."
        )
        return True

    stored_path = _store_text_blob(
        session["user_id"], text, suffix=filename.replace(".txt", "")
    )
    session["documents"].append(
        {
            "filename": filename,
            "text": text,
            "stored_path": stored_path,
            "warnings": [],
            "source_bytes": len(text.encode("utf-8")),
        }
    )
    session["total_characters"] = projected_total
    _set_session(context, session)

    await update.message.reply_text(
        f"📝 Captured inline text snippet ({len(text):,} chars). "
        f"Total documents: {len(session['documents'])}; total characters: {session['total_characters']:,}."
    )
    return True


# ---------------------------------------------------------------------------
# Core processing helpers


async def _process_import_documents(
    bot: "UnifiedWellnessBot",
    user_id: int,
    documents: Sequence[ImportDocument],
) -> dict:
    """Combine documents, extract structured profile data, and persist results."""
    if not documents:
        return {
            "total_documents": 0,
            "total_characters": 0,
            "stored_at": _now(),
            "documents": [],
            "extracted_profile": {},
            "user_summary": "No documents were processed.",
        }

    combined_text, total_characters = _combine_documents(documents)
    combined_path = _store_text_blob(user_id, combined_text, suffix="combined_import")

    extraction = await _extract_profile_insights(
        combined_text[:MAX_CONTEXT_CHARACTERS],
        total_characters,
        len(documents),
    )

    payload = {
        "stored_at": _now(),
        "total_documents": len(documents),
        "total_characters": total_characters,
        "combined_text_path": combined_path,
        "documents": [
            {
                "filename": doc.filename,
                "char_count": doc.char_count,
                "stored_path": doc.stored_path,
                "warnings": doc.warnings,
                "source_bytes": doc.source_bytes,
            }
            for doc in documents
        ],
        "extracted_profile": extraction or {},
    }
    payload["user_summary"] = _build_summary_line(payload)
    payload["preview"] = _truncated_preview(combined_text, SUMMARY_PREVIEW_CHARS)

    _persist_profile_import(user_id, payload)

    ingestion_meta = None
    try:
        ingestion_meta = await bot.ingest_profile_documents(
            user_id=user_id,
            combined_text=combined_text,
            combined_path=combined_path,
            source_documents=[doc.filename for doc in documents],
            total_characters=total_characters,
        )
    except Exception as exc:  # noqa: BLE001
        logger.error(
            "Profile import RAG ingestion failed for user %s: %s",
            user_id,
            exc,
            exc_info=True,
        )

    if ingestion_meta:
        payload["rag_ingestion"] = ingestion_meta

    # todo: Allow users to specify which imported documents should be discoverable via retrieval
    return payload


async def _extract_profile_insights(
    combined_text: str,
    total_characters: int,
    document_count: int,
) -> dict:
    """Call LLM to derive structured profile data from combined text."""
    if not combined_text.strip():
        return {}

    prompt = (
        "You are a personalization analyst. Extract concise user profile data from imported documents.\n\n"
        f"Document count: {document_count}\n"
        f"Total characters: {total_characters}\n\n"
        "Return JSON with these keys (omit any fields you cannot support with evidence):\n"
        "{\n"
        '  "personal_facts": {\n'
        '    "occupation": string | null,\n'
        '    "location": string | null,\n'
        '    "age_range": string | null,\n'
        '    "hobbies": [string],\n'
        '    "favorite_topics": [string]\n'
        "  },\n"
        '  "communication_preferences": {\n'
        '    "formality_level": 0.0-1.0 | null,\n'
        '    "verbosity": "brief" | "moderate" | "detailed" | null,\n'
        '    "tone": [string]\n'
        "  },\n"
        '  "wellness_context": {\n'
        '    "goals": [string],\n'
        '    "challenges": [string],\n'
        '    "mental_health_notes": string | null\n'
        "  },\n"
        '  "personality_highlights": [string],\n'
        '  "notable_quotes": [string]\n'
        "}\n\n"
        "Be concise. Use direct quotes only when they capture important identity details."
    )

    try:
        loop = asyncio.get_running_loop()
        result = await loop.run_in_executor(
            None,
            lambda: generate(
                prompt=f"{prompt}\n\nUSER DOCUMENTS:\n{combined_text}",
                format="json",
                options={"temperature": 0.1},
            ),
        )
    except Exception as exc:  # noqa: BLE001
        return {"error": f"LLM extraction failed: {exc}"}

    raw = (result or {}).get("text") or ""
    if not raw:
        return {}
    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        return {"raw_response": raw}


# ---------------------------------------------------------------------------
# Session management helpers


def _start_session(context: ContextTypes.DEFAULT_TYPE, user_id: int) -> None:
    user_data = context.user_data
    if user_data is None:
        return
    user_data[SESSION_KEY] = {
        "user_id": user_id,
        "documents": [],
        "total_characters": 0,
        "started_at": _now(),
        "inline_counter": 0,
    }


def _get_session(context: ContextTypes.DEFAULT_TYPE) -> Optional[dict]:
    user_data = context.user_data
    if user_data is None:
        return None
    session = user_data.get(SESSION_KEY)
    if session and session.get("documents") is not None:
        return session
    return None


def _set_session(context: ContextTypes.DEFAULT_TYPE, session: dict) -> None:
    user_data = context.user_data
    if user_data is None:
        return
    user_data[SESSION_KEY] = session


def _clear_session(context: ContextTypes.DEFAULT_TYPE) -> None:
    user_data = context.user_data
    if user_data is None:
        return
    user_data.pop(SESSION_KEY, None)


# ---------------------------------------------------------------------------
# Persistence and formatting helpers


def _persist_profile_import(user_id: int, content: Any) -> None:
    """Store import payload into profile_context."""
    if isinstance(content, str):
        serialized = content
    else:
        serialized = json.dumps(content, ensure_ascii=False)

    with db_rw() as conn:
        conn.execute(
            """
            INSERT INTO profile_context (user_id, key, value)
            VALUES (?, 'chatgpt_import', ?)
            ON CONFLICT(user_id, key)
            DO UPDATE SET value = excluded.value, updated_at = CURRENT_TIMESTAMP
            """,
            (user_id, serialized),
        )


def _format_user_summary(payload: dict) -> str:
    """Generate a conversational summary for user confirmation messages."""
    header = payload.get("user_summary") or "Import complete."
    extracted = payload.get("extracted_profile") or {}
    lines = [f"✅ {header}"]

    personal = extracted.get("personal_facts") or {}
    highlight_parts: list[str] = []
    for key in ("occupation", "location", "age_range"):
        value = personal.get(key)
        if value:
            highlight_parts.append(f"{key.replace('_', ' ').title()}: {value}")
    hobbies = personal.get("hobbies") or []
    if hobbies:
        highlight_parts.append(f"Hobbies: {', '.join(hobbies[:3])}")
    if highlight_parts:
        lines.append("• " + " | ".join(highlight_parts))

    wellness = extracted.get("wellness_context") or {}
    goals = wellness.get("goals") or []
    if goals:
        lines.append("• Goals captured: " + ", ".join(goals[:3]))
    challenges = wellness.get("challenges") or []
    if challenges:
        lines.append("• Challenges noted: " + ", ".join(challenges[:3]))

    notes = extracted.get("notable_quotes") or []
    if notes:
        lines.append("• Notable quote: " + notes[0][:180])

    return "\n".join(lines[:5])


def _build_summary_line(payload: dict) -> str:
    total_docs = payload.get("total_documents", 0)
    total_chars = payload.get("total_characters", 0)
    return f"Processed {total_docs} document(s) totaling {total_chars:,} characters."


def _combine_documents(documents: Sequence[ImportDocument]) -> Tuple[str, int]:
    combined: list[str] = []
    for doc in documents:
        header = f"FILENAME: {doc.filename}\n\n"
        combined.append(header + doc.text.strip())
    joined = "\n\n=== NEXT DOCUMENT ===\n\n".join(combined)
    return joined, sum(doc.char_count for doc in documents)


def _materialize_documents(raw_docs: Iterable[dict]) -> List[ImportDocument]:
    docs: list[ImportDocument] = []
    for raw_doc in raw_docs:
        docs.append(
            ImportDocument(
                filename=raw_doc.get("filename", "document"),
                text=raw_doc.get("text", ""),
                stored_path=raw_doc.get("stored_path"),
                warnings=raw_doc.get("warnings", []),
                source_bytes=raw_doc.get("source_bytes"),
            )
        )
    return docs


def _truncated_preview(text: str, limit: int) -> str:
    snippet = text.strip()
    if len(snippet) <= limit:
        return snippet
    return snippet[:limit].rstrip() + "…"


# ---------------------------------------------------------------------------
# Storage helpers


def _store_text_blob(user_id: int, text: str, suffix: str) -> str:
    base_dir = Path(user_dir(user_id)) / "derived" / "imports"
    base_dir.mkdir(parents=True, exist_ok=True)
    timestamp = operator_now().strftime("%Y%m%d_%H%M%S")
    safe_suffix = suffix.replace(" ", "_")
    filename = f"{timestamp}_{safe_suffix}.txt"
    target = base_dir / filename
    target.write_text(text, encoding="utf-8")
    return str(target.relative_to(user_dir(user_id)))


def _store_binary_document(user_id: int, original_name: str | None, data: bytes) -> str:
    base_dir = Path(user_dir(user_id)) / "media" / "documents"
    base_dir.mkdir(parents=True, exist_ok=True)
    timestamp = operator_now().strftime("%Y%m%d_%H%M%S")
    safe_name = _sanitize_filename(original_name or "document.bin")
    target = base_dir / f"{timestamp}_{safe_name}"
    target.write_bytes(data)
    return str(target.relative_to(user_dir(user_id)))


# ---------------------------------------------------------------------------
# Extraction helpers


def _extract_text_from_bytes(data: bytes, filename: str) -> Tuple[str, list[str]]:
    ext = Path(filename).suffix.lower()
    warnings: list[str] = []

    if ext in {".txt", ".md"}:
        try:
            text = data.decode("utf-8")
        except UnicodeDecodeError:
            text = data.decode("utf-8", errors="replace")
            warnings.append("Non-UTF8 characters were replaced.")
        return text, warnings

    if ext == ".json":
        try:
            content = json.loads(data.decode("utf-8"))
        except (json.JSONDecodeError, UnicodeDecodeError):
            warnings.append("JSON parsing failed; stored raw text.")
            return data.decode("utf-8", errors="replace"), warnings

        if isinstance(content, list):
            parts: list[str] = []
            for item in content:
                parts.extend(_extract_json_parts(item))
            return "\n\n".join(filter(None, parts)), warnings
        if isinstance(content, dict):
            return json.dumps(content, indent=2), warnings
        return str(content), warnings

    if ext == ".pdf":
        try:
            import PyPDF2  # type: ignore
        except ImportError:
            warnings.append("PyPDF2 not installed; stored binary PDF path only.")
            return "", warnings

        reader = PyPDF2.PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            try:
                pages.append(page.extract_text() or "")
            except Exception:  # noqa: BLE001
                pages.append("")
        text = "\n\n".join(pages)
        if not text.strip():
            warnings.append("PDF text extraction returned empty content.")
        return text, warnings

    if ext in {".docx", ".doc"}:
        try:
            import docx  # type: ignore
        except ImportError:
            warnings.append(
                "python-docx not installed; stored binary document path only."
            )
            return "", warnings

        document = docx.Document(io.BytesIO(data))
        text = "\n\n".join(paragraph.text for paragraph in document.paragraphs)
        if not text.strip():
            warnings.append("DOCX file contained no extractable text.")
        return text, warnings

    # Unknown file type
    warnings.append(
        f"Unsupported file type: {ext or 'unknown'}; stored binary data only."
    )
    return "", warnings


def _extract_json_parts(item: Any) -> List[str]:
    parts: list[str] = []
    if isinstance(item, dict):
        if "content" in item:
            content = item["content"]
            if isinstance(content, dict):
                inner = content.get("parts")
                if isinstance(inner, list):
                    parts.extend(
                        str(p) for p in inner if isinstance(p, (str, int, float))
                    )
            elif isinstance(content, list):
                parts.extend(
                    str(x) for x in content if isinstance(x, (str, int, float))
                )
            elif isinstance(content, str):
                parts.append(content)
        for value in item.values():
            parts.extend(_extract_json_parts(value))
    elif isinstance(item, list):
        for value in item:
            parts.extend(_extract_json_parts(value))
    elif isinstance(item, (str, int, float)):
        parts.append(str(item))
    return parts


# ---------------------------------------------------------------------------
# Utility helpers


def _sanitize_filename(name: str) -> str:
    candidate = Path(name).name
    return "".join(
        ch if ch.isalnum() or ch in {"-", "_", "."} else "_" for ch in candidate
    )


def _now() -> str:
    return operator_now().strftime("%Y-%m-%d %H:%M:%S")
